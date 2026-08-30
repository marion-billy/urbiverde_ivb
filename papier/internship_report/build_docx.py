"""Merge the internship-report markdown sections into a single formatted .docx.

Handles the markdown subset actually used in rapport_0..7:
- headings (#, ##, ###), hard-wrapped paragraphs (joined), markdown tables,
- encadré blockquotes -> shaded bordered boxes, brouillon notes -> dropped,
- [FIGURE : ...] markers -> embedded image if available else grey placeholder box,
- [TABLEAU : ...] markers -> table caption,
- inline **bold** and highlighted placeholders ([CHIFFRE], [À COMPLÉTER/VÉRIFIER/ADAPTER]).

Captions use Word SEQ fields + a Caption style, and the front matter holds auto fields:
a Sommaire (TOC), a Table des figures and a Liste des tableaux. In Word, select all and
press F9 (or right-click > Update fields) to populate them.

Run: python3 build_docx.py  ->  rapport_complet.docx
"""

from __future__ import annotations

import glob
import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE: str = os.path.dirname(os.path.abspath(__file__))
FIG: str = os.path.join(BASE, "figures")
PIPELINE: str = os.path.normpath(os.path.join(BASE, "..", "..", "data", "outputs", "pipeline.png"))

ENCADRE_BG: str = "EAF3EA"
PLACE_BG: str = "F4F4F4"

# Table des acronymes (front matter). Sigles réellement employés dans le rapport, développements
# standards. Ordre alphabétique. Un ajout au corps de texte reste défini à sa 1re occurrence : cette
# table est complémentaire, pas un substitut.
ACRONYMS: list[tuple[str, str]] = [
    ("EC", "Equivalent Connectivity (connectivité équivalente)"),
    ("ESA", "European Space Agency (Agence spatiale européenne)"),
    ("GBIF", "Global Biodiversity Information Facility"),
    ("IGN", "Institut national de l'information géographique et forestière"),
    ("INPN", "Inventaire national du patrimoine naturel"),
    ("MSPA", "Morphological Spatial Pattern Analysis (analyse morphologique des formes spatiales)"),
    ("OCS", "occupation du sol"),
    ("OSM", "OpenStreetMap"),
    ("PC", "Probability of Connectivity (probabilité de connectivité)"),
    ("PLU / PLUi", "plan local d'urbanisme (intercommunal)"),
    ("QGIS", "logiciel de système d'information géographique libre"),
    ("SCoT", "schéma de cohérence territoriale"),
    ("SIG", "système d'information géographique"),
    ("SRADDET", "schéma régional d'aménagement, de développement durable et d'égalité des territoires"),
    ("SRCE", "schéma régional de cohérence écologique"),
    ("TVB", "Trame verte et bleue"),
    ("UTM", "Universal Transverse Mercator (projection cartographique)"),
    ("VIIRS", "Visible Infrared Imaging Radiometer Suite (capteur d'imagerie nocturne)"),
]

PLACEHOLDER_RE = re.compile(r"\[(?:CHIFFRE|À (?:COMPLÉTER|VÉRIFIER|ADAPTER))[^\]]*\]")
TOKEN_RE = re.compile(r"(\*\*.+?\*\*|\[(?:CHIFFRE|À (?:COMPLÉTER|VÉRIFIER|ADAPTER))[^\]]*\])")


def figure_image(caption: str) -> str | None:
    """Map a figure caption to an available image path, or None for a placeholder."""
    # Règle générale : si la légende cite un fichier NAME.png (n'importe où, ex. "(figures/x.png)")
    # et qu'il existe, on l'utilise directement.
    m = re.search(r"([\w-]+\.png)", caption)
    if m and os.path.exists(os.path.join(FIG, m.group(1))):
        return os.path.join(FIG, m.group(1))
    c = caption.lower()
    if "arborescence" in c:
        return os.path.join(FIG, "arborescence_sorties.png")
    if "comparaison sur la rochelle" in c:
        return os.path.join(FIG, "comparaison_larochelle.png")
    if "gabriel" in c:
        return os.path.join(FIG, "gabriel_criterion.png")
    if "principe du chemin de moindre coût" in c:
        return os.path.join(FIG, "lcp_principe.png")
    if "familles de modèles de connectivité" in c:
        return os.path.join(FIG, "diniz_connectivity_models.png")
    if "tableau de bord interactif" in c:
        return os.path.join(FIG, "dashboard_capture.png")
    if "courbes de réponse à la distance" in c:
        return os.path.join(FIG, "sens_curves_d0.png")
    if "courbes de réponse au contraste" in c:
        return os.path.join(FIG, "sens_curves_contrast.png")
    if "part d'habitat connecté par territoire" in c:
        return os.path.join(FIG, "territoire_connectivite.png")
    if "réseau écologique" in c:
        return os.path.join(FIG, "reseau_ecologique.png")
    if "synoptique de la chaîne" in c:
        return os.path.join(FIG, "pipeline_chaine.png")
    if "segmentation morphologique" in c:
        return os.path.join(FIG, "mspa_example.png")
    if "complet de sorties" in c:
        return os.path.join(FIG, "sorties_perpignan.png")
    if "dispersion compar" in c:
        return os.path.join(FIG, "dispersion_comparee.png")
    if "occupation du sol" in c and "toulouse" in c:
        return os.path.join(FIG, "landcover_toulouse.png")
    if "localisation" in c:
        return os.path.join(FIG, "localisation_territoires_osm.png")
    if "gantt" in c:
        return os.path.join(FIG, "gantt_stage.png")
    if "côte à côte" in c:
        return os.path.join(FIG, "territoires_comparees.png")
    if "scénario de végétalisation" in c:
        return os.path.join(FIG, "scenario_local.png")
    if "poolés sur les cinq territoires" in c:
        return os.path.join(FIG, "gbif_ratios_pooled.png")
    if "par territoire et poolés" in c or "par territoire et agrégés" in c:
        return os.path.join(FIG, "gbif_recap_par_ville.png")
    if "fauvette à toulouse" in c or "oiseau de lisière à toulouse" in c:
        return os.path.join(FIG, "gbif_carte_Toulouse_fauvette.png")
    if "agrégés sur les cinq territoires" in c:
        return os.path.join(FIG, "gbif_ratios_pooled.png")
    if "sorties de la chaîne sur perpignan" in c:
        return os.path.join(FIG, "sorties_perpignan.png")
    if "salles" in c:
        return os.path.join(FIG, "comparaison_salles.png")
    if "cartes d'occurrences" in c and "oiseaux de lisière" in c:
        return os.path.join(FIG, "gbif_cartes_forest_edge_bird.png")
    if "cartes d'occurrences" in c and "mammifères arboricoles" in c:
        return os.path.join(FIG, "gbif_cartes_arboreal_mammal.png")
    if "cartes d'occurrences" in c and "mammifères terrestres" in c:
        return os.path.join(FIG, "gbif_cartes_ground_mammal.png")
    if "cartes d'occurrences" in c and "reptiles terrestres" in c:
        return os.path.join(FIG, "gbif_cartes_ground_reptile.png")
    if "tornade" in c and "hérisson" in c:
        return os.path.join(FIG, "tornado_ground_mammal_connected_habitat_pct.png")
    if "tornade" in c and "lézard" in c:
        return os.path.join(FIG, "tornado_ground_reptile_connected_habitat_pct.png")
    if "courbe de réponse" in c and "distance de dispersion" in c and "lézard" in c:
        return os.path.join(FIG, "response_d0_ground_reptile.png")
    if "courbe de réponse" in c and "distance de dispersion" in c and "hérisson" in c:
        return os.path.join(FIG, "response_d0_ground_mammal.png")
    if "courbe de réponse" in c and "contraste" in c and "lézard" in c:
        return os.path.join(FIG, "response_contrast_ground_reptile.png")
    if "courbe de réponse" in c and "contraste" in c and "hérisson" in c:
        return os.path.join(FIG, "response_contrast_ground_mammal.png")
    return None


def shade(el, hexcolor: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    el.append(shd)


def cell_background(cell, hexcolor: str) -> None:
    shade(cell._tc.get_or_add_tcPr(), hexcolor)


def box_borders(table, color: str = "9CB89C", size: str = "8") -> None:
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), size)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


def field(paragraph, instr: str, placeholder: str) -> None:
    """Insert a Word field (begin/instr/separate/placeholder/end) into a paragraph."""
    run = paragraph.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate")
    tt = OxmlElement("w:t"); tt.text = placeholder
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    for x in (fb, it, fs, tt, fe):
        run._r.append(x)


def add_runs(p, text: str) -> None:
    """Add inline runs: **bold** and highlighted placeholders."""
    for tok in TOKEN_RE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            p.add_run(tok[2:-2]).bold = True
        elif PLACEHOLDER_RE.fullmatch(tok):
            p.add_run(tok).font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            p.add_run(tok)


def add_caption(doc, label: str, num: str, text: str, center: bool = False):
    """Caption paragraph 'Label X.Y. text' with a hard-coded number (correct in any viewer)."""
    try:
        p = doc.add_paragraph(style="Caption")
    except KeyError:
        p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead = p.add_run(f"{label} {num}. ")
    lead.bold = True
    add_runs(p, text)
    return p


OBJET_RE = re.compile(r"\s*\(objet\s*:[^)]*\)")
# Working annotations of the bibliography, stripped at render: verification/correction notes and
# the "[URL ... à coller au rendu]" placeholders left on the legal-text entries.
WORK_ANNOT_RE = re.compile(r"\s*\[(?:non revérifiée|⚠|URL)[^\]]*\]")


def add_body_paragraph(doc, text: str) -> None:
    text = OBJET_RE.sub("", text)  # drop the "(objet : ...)" reading aid before rendering
    text = WORK_ANNOT_RE.sub("", text)  # drop working annotations ([non revérifiée], [⚠ ...], URL placeholders)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, text)


def add_encadre(doc, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    box_borders(table)
    cell = table.cell(0, 0)
    cell_background(cell, ENCADRE_BG)
    cell.paragraphs[0].text = ""
    paras: list[str] = []
    cur: list[str] = []
    for ln in lines:
        if ln.strip() == "":
            if cur:
                paras.append(" ".join(cur)); cur = []
        else:
            cur.append(ln.strip())
    if cur:
        paras.append(" ".join(cur))
    for i, para in enumerate(paras):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, para)
    doc.add_paragraph()


def classify_quote(content: list[str]) -> str:
    """Encadré (title line, blank, body) vs plain quote (e.g. the problématique) vs meta note."""
    firstne = next((ln for ln in content if ln.strip()), "").strip()
    if firstne.startswith(("Brouillon", "Front matter", "Travaux cités", "Références fournies",
                           "**Références encore", "Références encore")):
        return "drop"
    idx = next((i for i, ln in enumerate(content) if ln.strip()), None)
    if idx is not None and idx + 1 < len(content) and content[idx + 1].strip() == "" \
            and any(ln.strip() for ln in content[idx + 2:]):
        return "encadre"
    return "quote"


def add_quote(doc, content: list[str]) -> None:
    joined = " ".join(ln.strip() for ln in content if ln.strip())
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.6)
    p.paragraph_format.right_indent = Inches(0.6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, joined)
    doc.add_paragraph()


def add_figure(doc, caption: str, counters: dict) -> None:
    counters["fig"] += 1
    num = f"{counters['fig']}"
    img = figure_image(caption)
    if img and os.path.exists(img):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(img, width=Inches(6.0))
    else:
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        box_borders(table, color="BBBBBB", size="6")
        cell = table.cell(0, 0)
        cell_background(cell, PLACE_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("[Figure à insérer]")
        r.bold = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    add_caption(doc, "Figure", num, caption, center=True)
    doc.add_paragraph()


def add_equation(doc, num: str, latex: str) -> None:
    """Add a numbered display equation. The math is typeset via matplotlib mathtext into a PNG and
    embedded centred, with its number "(N)" to the right (mirrors the LaTeX equation/\\tag{N})."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    png = os.path.join(FIG, f"eq_{num}.png")
    fig = plt.figure(figsize=(0.1, 0.1))
    fig.text(0.5, 0.5, f"${latex}$", fontsize=18, ha="center", va="center")
    fig.savefig(png, dpi=150, bbox_inches="tight", pad_inches=0.06, transparent=True)
    plt.close(fig)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(png)
    r = p.add_run(f"  ({num})")
    r.font.size = Pt(11)
    doc.add_paragraph()


def add_table_caption(doc, caption: str, counters: dict) -> None:
    counters["tab"] += 1
    num = f"{counters['tab']}"
    add_caption(doc, "Tableau", num, caption)


def add_md_table(doc, rows: list[str]) -> None:
    parsed = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    header, body = parsed[0], parsed[2:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(header):
        c = table.rows[0].cells[j]
        cell_background(c, "D9E6D9")
        c.paragraphs[0].text = ""
        add_runs(c.paragraphs[0], h)
        for run in c.paragraphs[0].runs:
            run.bold = True
    for row in body:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            if j < len(cells):
                cells[j].paragraphs[0].text = ""
                add_runs(cells[j].paragraphs[0], val)
    doc.add_paragraph()


def is_table_block(block: list[str]) -> bool:
    return len(block) >= 2 and all("|" in ln for ln in block) and re.match(r"^\|?\s*-", block[1].strip()) is not None


def set_chapter(counters: dict, title: str) -> None:
    """Chapter number for figure/table numbering: parse a leading '2.' etc.; Annexes = last+1."""
    m = re.match(r"(\d+)\.", title.strip())
    if m:
        counters["chap"] = int(m.group(1)); counters["last"] = int(m.group(1))
    elif title.strip().lower().startswith("annexe"):
        counters["chap"] = counters.get("last", 0) + 1
    # Numérotation PLATE (continue sur tout le document) : pas de remise à zéro par chapitre.


def collect_captions(blocks: list[list[str]]):
    """First pass: chapter-aware numbers for the manual Table des figures / Liste des tableaux."""
    c = {"chap": 0, "fig": 0, "tab": 0, "last": 0}
    figs, tabs = [], []
    for block in blocks:
        first = block[0].lstrip()
        if first.startswith("#"):
            mm = re.match(r"(#+)\s+(.*)", first)
            if len(mm.group(1)) == 1:
                set_chapter(c, mm.group(2))
            continue
        joined = " ".join(l.strip() for l in block).strip()
        mf = re.match(r"\[FIGURE\s*:\s*(.*)\]$", joined)
        mt = re.match(r"\[TABLEAU\s*:\s*(.*)\]$", joined)
        if mf:
            c["fig"] += 1; figs.append((f"{c['fig']}", mf.group(1).strip()))
        elif mt:
            c["tab"] += 1; tabs.append((f"{c['tab']}", mt.group(1).strip()))
    return figs, tabs


def render_block(doc, block: list[str], counters: dict) -> None:
    first = block[0].lstrip()
    if first.startswith("#"):
        m = re.match(r"(#+)\s+(.*)", first)
        level = min(len(m.group(1)), 3)
        if level == 1:
            doc.add_page_break()
            set_chapter(counters, m.group(2))
        doc.add_heading(m.group(2).strip(), level=level)
        return
    if first.startswith(">"):
        content = [re.sub(r"^>\s?", "", ln) for ln in block]
        kind = classify_quote(content)
        if kind == "encadre":
            add_encadre(doc, content)
        elif kind == "quote":
            add_quote(doc, content)
        return
    if is_table_block(block):
        add_md_table(doc, block)
        return
    joined = " ".join(ln.strip() for ln in block).strip()
    mfig = re.match(r"\[FIGURE\s*:\s*(.*)\]$", joined)
    if mfig:
        add_figure(doc, mfig.group(1).strip(), counters)
        return
    mtab = re.match(r"\[TABLEAU\s*:\s*(.*)\]$", joined)
    if mtab:
        add_table_caption(doc, mtab.group(1).strip(), counters)
        return
    meq = re.match(r"\[EQUATION\s+(\d+)\s*:\s*(.*)\]$", joined)
    if meq:
        add_equation(doc, meq.group(1), meq.group(2).strip())
        return
    add_body_paragraph(doc, joined)


def parse_blocks(text: str) -> list[list[str]]:
    blocks: list[list[str]] = []
    cur: list[str] = []
    for line in text.splitlines():
        if line.strip() == "":
            if cur:
                blocks.append(cur); cur = []
        else:
            if line.lstrip().startswith("#") and cur:
                blocks.append(cur); cur = []
            cur.append(line)
    if cur:
        blocks.append(cur)
    return blocks


def _center(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _logo_cell(cell, name: str, height_in: float) -> None:
    path = os.path.join(BASE, "figures", name)
    p = _center(cell.paragraphs[0])
    if os.path.exists(path):
        p.add_run().add_picture(path, height=Inches(height_in))


def add_front_matter(doc, figs, tabs) -> None:
    # --- Page de garde : bandeau institutionnel ---
    banner = doc.add_table(rows=1, cols=3)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    _logo_cell(banner.rows[0].cells[0], "logo_univ_toulouse.png", 0.5)
    _logo_cell(banner.rows[0].cells[1], "logo_utjj.png", 0.4)
    _logo_cell(banner.rows[0].cells[2], "logo_inp_agro.png", 0.55)

    doc.add_paragraph()
    for txt, sz, bold in (("Master Géomatique SIGMA", 13, True),
                          ("ScIences Géomatiques en environneMent et Aménagement", 11, False),
                          ("Université de Toulouse", 10, False)):
        p = _center(doc.add_paragraph()); r = p.add_run(txt); r.font.size = Pt(sz); r.bold = bold

    for _ in range(2):
        doc.add_paragraph()
    p = _center(doc.add_paragraph()); r = p.add_run("RAPPORT DE MASTER 2"); r.font.size = Pt(12); r.bold = True
    doc.add_paragraph()
    p = _center(doc.add_paragraph())
    r = p.add_run("Conception d'un outil d'identification des continuités écologiques urbaines potentielles")
    r.bold = True; r.font.size = Pt(20)
    p = _center(doc.add_paragraph())
    r = p.add_run("à partir de l'observation de la Terre et de données ouvertes")
    r.font.size = Pt(14)

    for _ in range(3):
        doc.add_paragraph()
    p = _center(doc.add_paragraph()); r = p.add_run("Marion Billy"); r.bold = True; r.font.size = Pt(14)
    doc.add_paragraph()
    p = _center(doc.add_paragraph()); r = p.add_run("Structure d'accueil"); r.font.size = Pt(10)
    p = _center(doc.add_paragraph()); r = p.add_run("Murmuration"); r.bold = True; r.font.size = Pt(12)
    lg = os.path.join(BASE, "figures", "logo_murmuration.png")
    if os.path.exists(lg):
        _center(doc.add_paragraph()).add_run().add_picture(lg, height=Inches(0.45))

    # --- Encadrants et période : même page que la couverture ---
    doc.add_paragraph()
    for label in ("Tuteur de stage : Hugo Poupard",
                  "Enseignant-référent : Laurent Jégou",
                  "Période : du 16 février au 14 août 2026"):
        p = _center(doc.add_paragraph()); add_runs(p, label)

    # --- Sommaire (champ Word) + listes construites à la main ---
    doc.add_page_break()
    doc.add_heading("Sommaire", level=1)
    p = doc.add_paragraph()
    field(p, 'TOC \\o "1-3" \\h \\z \\u', "Ouvrir dans Word puis Ctrl+A et F9 pour générer le sommaire.")

    doc.add_heading("Table des figures", level=1)
    for num, text in figs:
        p = doc.add_paragraph()
        p.add_run(f"Figure {num} — ").bold = True
        add_runs(p, text.rstrip(". "))

    doc.add_heading("Liste des tableaux", level=1)
    for num, text in tabs:
        p = doc.add_paragraph()
        p.add_run(f"Tableau {num} — ").bold = True
        add_runs(p, text.rstrip(". "))

    # --- Table des acronymes (2 colonnes : sigle | signification) ---
    doc.add_heading("Table des acronymes", level=1)
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for sigle, sens in ACRONYMS:
        cells = tbl.add_row().cells
        cells[0].paragraphs[0].text = ""
        cells[0].paragraphs[0].add_run(sigle).bold = True
        cells[1].paragraphs[0].text = sens
    doc.add_paragraph()

    # --- Table des équations ---
    doc.add_heading("Table des équations", level=1)
    teq = doc.add_table(rows=0, cols=2)
    teq.style = "Table Grid"
    teq.alignment = WD_TABLE_ALIGNMENT.CENTER
    for num, desc in [("(1)", "Probabilité de déplacement entre deux taches"),
                      ("(2)", "Probability of Connectivity (PC)"),
                      ("(3)", "Budget de déplacement"),
                      ("(4)", "Surface équivalente connectée (EC)")]:
        cells = teq.add_row().cells
        cells[0].paragraphs[0].add_run(num).bold = True
        cells[1].paragraphs[0].text = desc
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    # Ask Word to update SEQ/TOC fields on open, so Figure/Tableau numbers and the lists fill in
    # without a manual Ctrl+A then F9.
    upd = OxmlElement("w:updateFields"); upd.set(qn("w:val"), "true")
    doc.settings.element.append(upd)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # A more professional layout: margins, coloured headings, footer page numbers.
    sec = doc.sections[0]
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(1.0))
    for lvl, sz in ((1, 16), (2, 13), (3, 11.5)):
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = "Times New Roman"
        h.font.size = Pt(sz)
        h.font.color.rgb = RGBColor(0x12, 0x3F, 0x3A)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field(fp, " PAGE ", "1")

    files = sorted(glob.glob(os.path.join(BASE, "rapport_*.md")))
    blocks_all = []
    for fp in files:
        blocks_all += parse_blocks(open(fp, encoding="utf-8").read())

    figs, tabs = collect_captions(blocks_all)          # pass 1: chapter-aware numbers
    add_front_matter(doc, figs, tabs)                  # title page + manual lists

    counters = {"chap": 0, "fig": 0, "tab": 0, "last": 0}
    for block in blocks_all:                            # pass 2: render body
        render_block(doc, block, counters)

    out = os.path.join(BASE, "rapport_complet.docx")
    doc.save(out)
    print("DOCX écrit :", out)
    print(f"  figures : {len(figs)} ; tableaux : {len(tabs)} ; sections : {len(files)}")


if __name__ == "__main__":
    main()
